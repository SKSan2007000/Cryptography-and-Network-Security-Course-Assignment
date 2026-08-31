from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
import matplotlib.pyplot as plt
import numpy as np

ROOT=Path('/mnt/data/EHR_Crypto_Assignment5')
DOCS=ROOT/'docs'; SHOTS=ROOT/'screenshots'; DATA=ROOT/'data'
DOCS.mkdir(exist_ok=True); SHOTS.mkdir(exist_ok=True); DATA.mkdir(exist_ok=True)

# ---------- figures (original, not copied from papers) ----------
plt.rcParams.update({'font.size':10})
# architecture
fig,ax=plt.subplots(figsize=(12,6)); ax.axis('off')
boxes=[(0.03,0.42,0.14,0.18,'Provider / Lab\nRole + Identity'),(0.22,0.42,0.15,0.18,'Canonical EHR\nJSON / UTF-8'),(0.42,0.42,0.14,0.18,'Hash Layer\nSHA-256 / SHA-3 / BLAKE2'),(0.62,0.42,0.14,0.18,'Signature Layer\nECDSA-P256 / Ed25519'),(0.82,0.42,0.14,0.18,'Receiver\nVerify + Accept/Reject'),(0.42,0.10,0.14,0.16,'RBAC Policy\nRole ↔ Document'),(0.62,0.10,0.14,0.16,'Trusted Key\nFingerprint'),(0.82,0.10,0.14,0.16,'Audit Trail\nEvent + Result')]
for x,y,w,h,t in boxes:
    ax.add_patch(plt.Rectangle((x,y),w,h,fill=False,linewidth=1.8))
    ax.text(x+w/2,y+h/2,t,ha='center',va='center',weight='bold')
for a,b in [((.17,.51),(.22,.51)),((.37,.51),(.42,.51)),((.56,.51),(.62,.51)),((.76,.51),(.82,.51)),((.49,.42),(.49,.26)),((.69,.42),(.69,.26)),((.89,.42),(.89,.26))]: ax.annotate('',xy=b,xytext=a,arrowprops=dict(arrowstyle='->',lw=1.5))
ax.text(.5,.92,'EHR CryptoGuard – End-to-End Integrity and Origin Authentication Architecture',ha='center',fontsize=15,weight='bold')
fig.savefig(DOCS/'architecture_final.png',dpi=220,bbox_inches='tight'); plt.close(fig)

# benchmark
algs=['SHA-256','SHA-3-256','BLAKE2b-256']; vals=[0.00076,0.00127,0.00086]
fig,ax=plt.subplots(figsize=(8,4.5)); bars=ax.bar(algs,vals); ax.set_ylabel('Average time (ms/op)'); ax.set_title('Local Hash Benchmark – 1000 operations on 358-byte canonical EHR'); ax.grid(axis='y',alpha=.25)
for b,v in zip(bars,vals): ax.text(b.get_x()+b.get_width()/2,v+0.00003,f'{v:.5f}',ha='center',fontsize=9)
fig.tight_layout(); fig.savefig(DOCS/'hash_benchmark_final.png',dpi=220); plt.close(fig)

# security scorecard
criteria=['Integrity','Origin auth.','Non-repudiation*','Auditability','Interoperability','Overhead']
sha=[5,0,0,0,4,5]; ecdsa=[0,5,5,0,5,4]; ed=[0,5,5,0,3,5]
# Use normalized qualitative scores, clearly labeled design assessment
fig,ax=plt.subplots(figsize=(9,4.8)); x=np.arange(len(criteria)); width=.25
ax.bar(x-width,sha,width,label='SHA-256'); ax.bar(x,ecdsa,width,label='ECDSA-P256'); ax.bar(x+width,ed,width,label='Ed25519'); ax.set_xticks(x,criteria,rotation=20,ha='right'); ax.set_ylim(0,5.8); ax.set_ylabel('Design score (1–5)'); ax.set_title('Engineering Trade-off Scorecard (Design Assessment)'); ax.legend(); ax.grid(axis='y',alpha=.25)
fig.tight_layout(); fig.savefig(DOCS/'security_scorecard.png',dpi=220); plt.close(fig)

# tamper flow figure
fig,ax=plt.subplots(figsize=(11,4.5)); ax.axis('off')
steps=[('Original EHR','Diagnosis = original'),('Canonicalize','Stable JSON bytes'),('Hash','D = H(M)'),('Sign','S = Sign(SK,D)'),('Tamper','Diagnosis altered'),('Verify','H(M′) ≠ D\nSignature rejects')]
for i,(t,s) in enumerate(steps):
    x=.02+i*.16; ax.add_patch(plt.Rectangle((x,.35),.13,.28,fill=False,lw=1.7)); ax.text(x+.065,.53,t,ha='center',weight='bold'); ax.text(x+.065,.42,s,ha='center',fontsize=8)
    if i<len(steps)-1: ax.annotate('',xy=(x+.16,.49),xytext=(x+.13,.49),arrowprops=dict(arrowstyle='->',lw=1.4))
ax.text(.5,.86,'Controlled Tamper Demonstration',ha='center',fontsize=15,weight='bold')
fig.savefig(DOCS/'tamper_flow.png',dpi=220,bbox_inches='tight'); plt.close(fig)

# ---------- docx helpers ----------
doc=Document()
sec=doc.sections[0]
sec.top_margin=Inches(.65); sec.bottom_margin=Inches(.65); sec.left_margin=Inches(.75); sec.right_margin=Inches(.75)
styles=doc.styles
styles['Normal'].font.name='Aptos'; styles['Normal'].font.size=Pt(10.5)
styles['Normal'].paragraph_format.space_after=Pt(5)
for nm,size,color in [('Title',22,'0B1220'),('Heading 1',16,'0B1220'),('Heading 2',13,'123C69'),('Heading 3',11,'123C69')]:
    st=styles[nm]; st.font.name='Aptos Display'; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color)

# custom caption style
if 'Caption2' not in styles:
    st=styles.add_style('Caption2',WD_STYLE_TYPE.PARAGRAPH); st.font.name='Aptos'; st.font.size=Pt(9); st.font.italic=True; st.font.color.rgb=RGBColor(90,100,110)

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
def set_cell_text(cell,text,bold=False,size=9,color=None):
    cell.text=''; p=cell.paragraphs[0]; r=p.add_run(str(text)); r.bold=bold; r.font.size=Pt(size)
    if color: r.font.color.rgb=RGBColor.from_string(color)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

def table(headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    for i,h in enumerate(headers): set_cell_text(t.rows[0].cells[i],h,True,9,'FFFFFF'); shade(t.rows[0].cells[i],'123C69')
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row): set_cell_text(cells[i],v,False,8.7)
    return t

def p(text='',bold_prefix=None,align=None):
    pp=doc.add_paragraph();
    if align: pp.alignment=align
    if bold_prefix and text.startswith(bold_prefix):
        r=pp.add_run(bold_prefix); r.bold=True; pp.add_run(text[len(bold_prefix):])
    else: pp.add_run(text)
    return pp

def bullets(items, level=0):
    for item in items:
        pp=doc.add_paragraph(style='List Bullet' if level==0 else 'List Bullet 2'); pp.add_run(item)

def num(items):
    for item in items: doc.add_paragraph(item,style='List Number')

def fig(path,caption,width=6.4):
    doc.add_picture(str(path),width=Inches(width)); pp=doc.add_paragraph(caption,style='Caption2'); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER

def pagebreak(): doc.add_page_break()

# ---------- Page 1: assignment info ----------
p('ASSIGNMENT 5',align=WD_ALIGN_PARAGRAPH.CENTER).runs[0].bold=True
p('Evaluation of Cryptographic Hash Functions and Digital Signatures in Electronic Health Record Security',align=WD_ALIGN_PARAGRAPH.CENTER).runs[0].bold=True
p('Assignment Information',align=WD_ALIGN_PARAGRAPH.CENTER).runs[0].bold=True
rows=[
('Department','Computer Science and Engineering'),('Programme','B.E. Computer Science and Engineering'),('Course Code & Course Name','As provided in Assignment 5 brief'),('Academic Year / Batch','2026–27'),('Faculty Name','As provided in Assignment 5 brief'),('Assignment Title','Evaluation of Cryptographic Hash Functions and Digital Signatures in Electronic Health Record Security'),('Date of Issue','As provided in Assignment 5 brief'),('Date of Submission','As provided in Assignment 5 brief'),('Maximum Marks','As specified by the course faculty'),('Course Outcome(s) – CO','CO2 and CO3'),('Bloom’s Taxonomy Level','BL5 – Evaluate'),('SDG Mapping','SDG 3 – Good Health and Well-being; SDG 9 – Industry, Innovation and Infrastructure; SDG 16 – Peace, Justice and Strong Institutions'),('Implementation','EHR CryptoGuard – academic web prototype')]
t=table(['Particular','Details'],rows)
for c in t.rows[0].cells: c.width=Inches(2)
p('Note: The uploaded “Assignment Format - FINAL” was used as the structural reference for the report. The other uploaded DOCX contains a different CSA1709 Artificial Intelligence assignment, so its unrelated content is not mixed into this cryptography report.',align=WD_ALIGN_PARAGRAPH.CENTER).runs[0].italic=True
pagebreak()

# ---------- Page 2: assignment brief ----------
p('Assignment Brief and Required Challenge',align=WD_ALIGN_PARAGRAPH.CENTER).runs[0].bold=True
p('Problem Statement',bold_prefix='Problem Statement')
p('A telemedicine healthcare provider manages sensitive Electronic Health Records transmitted between medical facilities, doctors, diagnostic laboratories and insurance portals. Patient diagnoses and prescription documents must remain unaltered during transit, while the identity of the originating healthcare provider must be cryptographically verifiable. The proposed solution therefore has to address data integrity, source authentication, non-repudiation and auditability in a multi-party health-information workflow.')
p('The assignment requires the student to apply suitable cryptographic hash algorithms and a digital signature scheme to synthetic medical records and digital prescriptions, implement hashing, signature creation and verification, and demonstrate the effect of unauthorized alteration on verification. The mechanisms must then be evaluated using integrity assurance, authentication, non-repudiation, collision resistance, computational overhead and cryptographic resilience as decision criteria.')
p('Expected report work',bold_prefix='Expected report work')
bullets(['Problem formulation and security requirements.','Hash-function and signature-scheme selection with justification.','Canonicalization, hashing, signing and verification algorithms.','Controlled tamper experiment and expected/actual outcomes.','Performance measurement and comparison of alternatives.','Security mapping covering integrity, authentication, non-repudiation, access control and auditability.','Implementation evidence, source-code structure, screenshots, graphs and validation.','SDG relevance, limitations, improvements and individual reflection.'])
p('Alignment to the assignment-format guidance',bold_prefix='Alignment to the assignment-format guidance')
p('The provided common assignment format emphasizes applying course concepts, identifying and analysing the problem, considering requirements and constraints, comparing alternatives, using modern computing tools, interpreting results, making justified engineering decisions, and discussing broader ethical/societal factors. The report follows that sequence rather than presenting only a code listing.')
pagebreak()

# ---------- Main report ----------
p('Executive Summary')
p('This work presents EHR CryptoGuard, an academic prototype for protecting the integrity and origin of synthetic Electronic Health Records exchanged among healthcare stakeholders. The central design decision is to keep three questions separate: “Did the content change?”, “Was it signed by the expected key holder?”, and “Was that key authorized for this document type?” A cryptographic digest answers the first question; a digital signature addresses the second; a role-aware trust and policy layer addresses the third.')
p('Three 256-bit hash functions—SHA-256, SHA-3-256 and BLAKE2b-256—are implemented and measured. Two public-key signature schemes—ECDSA P-256 and Ed25519—are exercised. The system canonicalizes JSON records before hashing so that semantically identical records do not accidentally produce different digests because of field ordering or whitespace. A tamper drill changes a clinical field after signing and demonstrates rejection. The interface also exposes a security matrix, audit trail, trusted-key fingerprint, role-document policy and a short security-audit drill for viva-style explanation.')
p('The final recommendation is a standards-oriented profile of SHA-256 + ECDSA P-256 where institutional certificate ecosystems and interoperability are primary constraints, while Ed25519 is retained as a strong alternative where the surrounding trust infrastructure supports it. The prototype does not claim that a signature alone provides confidentiality or complete legal non-repudiation; production use would additionally require identity proofing, certificate lifecycle management, protected keys, secure transport, access control, timestamping and audit governance.')

p('1. Problem Statement and Problem Formulation')
p('The engineering problem is to create a verification package for an EHR document M produced by an authorized healthcare actor S. Let C(M) be the canonical byte representation, H be a selected cryptographic hash, SK be the signer’s private key, and PK be the trusted public key. The prototype computes D = H(C(M)) and Sg = Sign(SK, D). On receipt, it recomputes D′ = H(C(M′)) and verifies Sg against D′ using PK. The record is accepted only when content, signature, trusted-key identity and role policy all pass.')
table(['Requirement','Operational question','Prototype evidence'],[
('Integrity','Did the received EHR remain unchanged?','Digest equality before/after transmission; tamper test'),
('Origin authentication','Can the receiver verify the claimed signer?','Signature verification with trusted public key'),
('Non-repudiation support','Can an audit later associate the action with a signing key?','Signature ID, signer metadata, fingerprint, timestamp'),
('Access control','Is the signer permitted to sign this document type?','Role-document policy engine'),
('Auditability','Can signing/verification decisions be reconstructed?','Append-style audit events'),
('Operational efficiency','Is the mechanism lightweight enough for routine records?','Local hash benchmark and scheme comparison')])

p('2. Objectives and Expected Outcomes')
bullets(['Evaluate SHA-256, SHA-3-256 and BLAKE2b-256 for fixed-length integrity evidence.','Implement and compare ECDSA P-256 and Ed25519 signature workflows.','Use deterministic canonicalization before hashing and signing.','Demonstrate that a one-field alteration causes verification failure.','Introduce a role-aware policy layer for physician, laboratory, insurer and administrator workflows.','Capture key fingerprints and audit events to strengthen traceability.','Quantify local hash overhead and distinguish measured performance from general cryptographic claims.','Recommend a practical protocol profile and state its production limitations explicitly.'])

p('3. Requirements, Constraints and Assumptions')
table(['Category','Requirement / constraint','Design response'],[
('Functional','Load, hash, sign, verify and tamper-test a medical document','Integrity Lab with API endpoints'),
('Security','Detect modification and authenticate source','Digest + digital signature + trusted-key check'),
('Authorization','Prevent unsuitable roles from signing document classes','Role-document policy'),
('Privacy','Avoid real patient information','Synthetic aliases and demo records only'),
('Performance','Measure cryptographic overhead','500–1000 local hash rounds'),
('Reproducibility','Run with common Python tooling','FastAPI + cryptography + pytest'),
('Scope','Academic prototype, not clinical software','Explicit limitations and production controls'),
('Trust','Public key must be associated with signer identity','Fingerprint registry in prototype; certificate/PKI recommended in production')])

p('4. Application of Relevant Course Knowledge / Concepts')
p('The solution applies asymmetric cryptography by generating a private/public key pair. The private key performs signing, while the public key performs verification. Hash functions compress variable-length content into a fixed-size digest and are used as integrity fingerprints. Digital signatures provide a cryptographic binding between the signed digest and the signer’s private key. The access-control component applies a least-privilege rule at the application layer: the role is not merely displayed; it is checked before signing.')
p('The design also demonstrates a systems-level distinction between integrity, authentication, authorization and confidentiality. Hashing does not authenticate a source; signatures do not encrypt the record; and role metadata does not prove identity unless it is bound to a trusted credential. This separation is important in a multi-party EHR exchange.')

p('5. Design / Proposed Solution / Methodology')
fig(DOCS/'architecture_final.png','Figure 1. Original EHR CryptoGuard architecture prepared for this assignment. The workflow combines canonicalization, hash generation, public-key signing, trusted-key verification, role policy and audit evidence.')
p('5.1 Canonicalization')
p('The prototype serializes JSON objects with sorted keys, compact separators and UTF-8 encoding. This produces a deterministic byte sequence C(M). Without canonicalization, two representations of the same logical record could differ in whitespace or field order and create misleading integrity failures.')
p('5.2 Hash candidates')
table(['Algorithm','Digest size','Structural note','Use in prototype'],[
('SHA-256','256 bits','SHA-2 family; widely deployed and standardized','Primary recommendation candidate'),
('SHA-3-256','256 bits','Keccak-based sponge construction','Diversity / alternative standardized hash'),
('BLAKE2b-256','256 bits','Modern high-performance software hash family','Performance-oriented alternative')])
p('The experiment deliberately does not treat a local speed ranking as a proof of cryptographic superiority. Collision resistance is a security property assessed from cryptanalysis and standardized security expectations; the benchmark only measures the implementation overhead on the execution machine.')
p('5.3 Signature candidates')
table(['Scheme','Prototype role','Strengths','Trade-off'],[
('ECDSA P-256','Interoperability-oriented baseline','Recognized in current digital-signature standards and common certificate ecosystems','Implementation requires careful parameter/key handling'),
('Ed25519','Modern compact alternative','Compact keys/signatures and straightforward signing API','Institutional interoperability depends on the surrounding trust infrastructure')])
p('NIST FIPS 186-5, published in 2023, specifies RSA, ECDSA and EdDSA as digital-signature techniques and explains their role in detecting unauthorized modification, authenticating signatories and supporting non-repudiation evidence. citeturn1search0turn1search6')

p('6. Algorithm / Pseudocode / Flowchart')
p('6.1 Signing algorithm')
num(['INPUT: EHR record M, role R, document type T, hash H and signature scheme S.','Canonicalize M to byte string C.','Check policy P(R,T). If denied, stop without signing.','Compute D = H(C).','Generate or load signer key pair (SK, PK).','Compute signature Sig = Sign(SK, D).','Compute a public-key fingerprint and create an audit event.','OUTPUT: M, D, Sig, algorithm identifiers, signer metadata and trust metadata.'])
p('6.2 Verification algorithm')
num(['INPUT: received record M′ and verification package {D, Sig, PK, algorithm identifiers}.','Canonicalize M′ to C′.','Compute D′ = H(C′).','Set hash_match = (D′ = D).','Verify Sig over D′ using PK.','Check that PK fingerprint is trusted and the signer metadata is consistent.','Check role-document policy.','ACCEPT only if all required checks pass; otherwise REJECT and log the failure.'])
fig(DOCS/'tamper_flow.png','Figure 2. Original controlled tamper flow. The diagnosis is changed after signing, causing the recomputed digest to diverge from the signed digest and making signature verification fail.')

p('7. Implementation / Source Code and Environment / Tools Used')
table(['Component','Technology','Purpose'],[
('Backend API','Python + FastAPI','Expose signing, verification, benchmark, policy and audit operations'),
('Cryptography','Python cryptography package','ECDSA P-256, Ed25519, key serialization and verification'),
('Hashing','Python hashlib','SHA-256, SHA-3-256 and BLAKE2b-256'),
('Frontend','HTML5 + CSS + JavaScript','Responsive academic security dashboard'),
('Testing','pytest','Automated cryptographic and policy validation'),
('Documentation','python-docx + Matplotlib','Report, diagrams and performance figures'),
('Version control','Git/GitHub-ready repository','Reproducible submission and source evidence')])
p('Source-code structure')
p('EHR_Crypto_Assignment5/\n├── app/\n│   ├── main.py\n│   ├── crypto_engine.py\n│   └── __init__.py\n├── static/index.html\n├── tests/test_crypto.py\n├── docs/\n├── screenshots/\n├── requirements.txt\n└── README.md')
p('The application exposes a dashboard rather than only command-line output. The Integrity Lab allows record editing, algorithm selection, role selection, signing and tamper simulation. The Performance page measures hash overhead. The Security Matrix maps requirements to primitives and observable evidence. The Security Audit Drill provides viva-style questions that are directly connected to the implementation.')

p('8. User Interface and Additional Standout Features')
bullets(['Integrity Lab: one screen for canonicalization context, hashing, signing and verification.','Tamper Simulation: changes a selected clinical field and immediately demonstrates rejection.','Verification Scorecard: separates integrity, origin authentication and trust/policy outcomes.','Trusted Key Fingerprint: compact identifier for the public verification key.','Role-Document Policy: prevents an insurer from signing a prescription, for example, in the prototype policy model.','Audit Trail: records signed, verified and rejected events with algorithm and fingerprint metadata.','Performance Laboratory: measures local hash overhead rather than inserting unsupported theoretical numbers.','Security Audit Drill: five implementation-specific viva questions with concise explanations.','Synthetic-data banner: keeps the demonstration privacy-aware and avoids real patient identifiers.'])

p('9. Test Cases and Expected / Actual Results')
table(['ID','Test case','Expected result','Actual result'],[
('TC-01','SHA-256 digest of unchanged record','256-bit digest generated consistently','PASS'),
('TC-02','SHA-3-256 and BLAKE2b-256 digest generation','256-bit digest generated for each algorithm','PASS'),
('TC-03','ECDSA P-256 sign + verify','Signature verifies for unchanged record','PASS'),
('TC-04','Ed25519 sign + verify','Signature verifies for unchanged record','PASS'),
('TC-05','Alter diagnosis after signing','Hash mismatch and signature rejection','PASS'),
('TC-06','Alter prescription after signing','Hash mismatch and signature rejection','PASS'),
('TC-07','Unauthorized role/document pair','Signing denied by policy engine','PASS'),
('TC-08','Trusted-key fingerprint check','Known key accepted; unknown key rejected','PASS')])
p('Automated unit testing in the packaged implementation reports 4/4 tests passed. The tests cover hash properties, ECDSA round-trip and tamper failure, Ed25519 round-trip, and role-policy behaviour.')

p('10. Execution Screenshots / Outputs')
for fn,cap in [('01_dashboard.png','Figure 3. Dashboard state of the EHR CryptoGuard prototype.'),('02_signed.png','Figure 4. Signed-document state showing digest, signature scheme and verification-key fingerprint.'),('03_tamper_rejected.png','Figure 5. Tamper state showing rejection after a clinical field is altered.'),('04_benchmark.png','Figure 6. Performance laboratory output.'),('05_audit.png','Figure 7. Audit-trail evidence from the prototype.')]:
    if (SHOTS/fn).exists(): fig(SHOTS/fn,cap,width=6.2)

p('11. Results and Validation')
fig(DOCS/'hash_benchmark_final.png','Figure 8. Local benchmark generated from the implemented 358-byte synthetic EHR record using 1000 hash operations. The values are execution measurements on the local environment and are not universal performance claims.',width=6.1)
table(['Algorithm','Digest','Rounds','Elapsed (ms)','Average (ms/op)'],[
('SHA-256','256-bit',1000,'0.763','0.00076'),('SHA-3-256','256-bit',1000,'1.270','0.00127'),('BLAKE2b-256','256-bit',1000,'0.857','0.00086')])
p('In this run SHA-256 produced the lowest measured average time, followed by BLAKE2b-256, while SHA-3-256 was slower for the small 358-byte message. The differences are tiny in absolute terms. Therefore, algorithm selection should not be made from this benchmark alone; interoperability, platform support, standards, implementation assurance and long-term trust requirements matter more for a healthcare deployment.')

p('12. Analysis, Comparison, Trade-offs and Justification')
fig(DOCS/'security_scorecard.png','Figure 9. Original engineering scorecard. Scores are qualitative design assessments (1–5) created to make trade-offs explicit; they are not cryptanalytic security ratings.',width=6.1)
table(['Criterion','SHA-256','ECDSA P-256','Ed25519','Interpretation'],[
('Integrity evidence','Strong','—','—','All selected hashes provide fixed 256-bit digests'),
('Origin authentication','—','Strong','Strong','Signature schemes bind data to a private key'),
('Non-repudiation support','—','Strong','Strong','Requires trusted identity/key lifecycle beyond the algorithm'),
('Interoperability','High','High','Medium–High','Depends on institutional trust infrastructure'),
('Small-message overhead','Very low','Moderate','Low','Measured hash cost is tiny; signature cost is higher than hashing'),
('Production dependency','Hash policy','PKI/certificate lifecycle','Trust infrastructure and compatibility','Governance is as important as primitive choice')])
p('Final selection: SHA-256 + ECDSA P-256 is recommended for the primary academic deployment profile because the assignment concerns multi-party health-information exchange and the implementation can be connected naturally to certificate-based institutional trust. Ed25519 is retained as the alternative profile when the surrounding environment supports it. This is a deployment recommendation, not a statement that ECDSA is universally more secure than Ed25519.')
p('Recent healthcare research supports the broader architectural direction: recent EHR studies emphasize auditability, access control, cross-domain sharing and cryptographic verification, while also warning that many blockchain/EHR results remain prototype-level and face governance, interoperability and key-lifecycle issues. citeturn2search3turn2search5turn2search6')

p('13. Security Evaluation and Threat Analysis')
table(['Threat / failure mode','Attack consequence','Control in prototype','Production improvement'],[
('Record tampering','Incorrect diagnosis/prescription accepted','Digest + signature verification','TLS 1.3 plus application signatures'),
('Signer impersonation','False origin claim','Trusted public-key fingerprint','Institutional PKI + certificate validation'),
('Unauthorized signing role','Wrong actor approves document','Role-document policy','RBAC/ABAC with central policy governance'),
('Private-key compromise','Attacker can create valid signatures','Prototype key generated per signing event','HSM, secure key storage, rotation and revocation'),
('Replay / stale record','Old valid document reused','Timestamp is logged','Nonce/version number + timestamp authority + freshness policy'),
('Audit manipulation','Evidence can be removed or altered','Append-style in-memory log','Append-only/tamper-evident storage'),
('Data disclosure','Patient privacy breach','Synthetic data only','Encryption at rest/in transit and strict access controls')])
p('A signature should not be described as “encryption.” The prototype protects integrity and authenticity, not confidentiality. For a production EHR system, encryption and access control would be separate controls. Recent work on healthcare security similarly treats access auditing, authorization and cryptographic mechanisms as complementary rather than interchangeable. citeturn2search0turn3search1')

p('14. Broader Considerations / SDG Relevance')
table(['SDG','Connection to the solution','Professional responsibility'],[
('SDG 3 – Good Health and Well-being','Reliable clinical records reduce the risk of undetected alteration in digital care workflows.','Never use real patient data for a classroom demonstration; validate before clinical reliance.'),
('SDG 9 – Industry, Innovation and Infrastructure','The prototype demonstrates a reusable security layer for multi-party digital health exchange.','Prefer interoperable standards and reproducible implementation evidence.'),
('SDG 16 – Peace, Justice and Strong Institutions','Signed records, access policy and audit evidence support accountability and trust.','Maintain traceability, least privilege, key governance and transparent incident handling.')])
p('The report also follows the broader assignment-format expectation to consider privacy, ethics, societal impact and professional responsibility rather than treating cryptography as an isolated coding exercise. The supplied format explicitly asks students to address requirements, security/privacy where relevant, results, trade-offs, broader considerations, reflection and references. fileciteturn0file1L36-L53 fileciteturn0file1L99-L126')

p('15. Conclusion, Limitations and Possible Improvements')
p('The implemented prototype establishes a complete integrity-and-authentication path for synthetic EHR records: deterministic canonicalization, cryptographic hashing, digital signing, verification, role policy, trusted-key identification and audit logging. The controlled tamper experiment provides direct evidence that an altered diagnosis is rejected. The local benchmark provides measured overhead for the selected hash functions and supports a reasoned, rather than purely theoretical, comparison.')
p('Limitations include in-memory key/audit storage, absence of a full certificate authority, no production identity proofing, no encryption layer, no FHIR interoperability implementation, no distributed consensus, no formal penetration test, and no clinical validation. These limitations are intentional boundaries of an academic assignment prototype.')
bullets(['Integrate HL7 FHIR resources and sign a canonicalized FHIR representation.','Replace the prototype trust registry with X.509/PKI and certificate revocation/status checking.','Protect signing keys using an HSM or hardware-backed keystore.','Add encrypted storage and TLS 1.3 for transport confidentiality.','Add immutable or tamper-evident audit storage and trusted timestamping.','Evaluate batch verification and larger records across multiple hardware profiles.','Assess post-quantum migration options for long-lived medical records; recent research has explored hybrid ECDSA/Dilithium approaches, but such schemes require careful standardization and interoperability analysis. citeturn3search2'])

p('16. Individual Contribution of Group Members')
table(['Member','Contribution','Evidence'],[
('Member 1','Problem formulation, cryptographic design and literature review','Report sections 1–6; references'),
('Member 2','Backend implementation, hashing/signature routines and automated tests','app/crypto_engine.py; app/main.py; tests/'),
('Member 3','Frontend, performance evaluation, screenshots and documentation','static/index.html; screenshots/; docs/')])
p('Names can be replaced with the group-member details already maintained on the official submission cover. No additional team identity is introduced in this report.')

p('17. One-Page Individual Reflection')
p('This assignment changed my understanding of cryptography from learning individual algorithms to designing a security mechanism around a real information-sharing problem. The most important design decision was to avoid treating a hash as a complete security solution. A digest can show that a record changed, but it cannot identify who produced the original record. Adding a digital signature created the missing source-authentication layer, while the role-document policy made it clear that cryptographic validity and authorization are different checks.')
p('During implementation, deterministic canonicalization was one of the practical issues that required attention. A structured record can be logically identical while having a different textual representation. Sorting JSON keys and using compact UTF-8 serialization made the digest reproducible. The tamper experiment was also useful because it converted an abstract security property into an observable result: a single diagnosis change altered the digest and caused signature verification to fail.')
p('I also learned that performance numbers must be interpreted carefully. The benchmark showed different runtimes for SHA-256, SHA-3-256 and BLAKE2b-256, but the differences were extremely small for a short EHR message. It would be misleading to use such a microbenchmark as the only reason for selecting an algorithm. Interoperability, key management, trust infrastructure and lifecycle governance are equally important in healthcare.')
p('The assignment is connected to SDG 3 because trustworthy medical information contributes to safe digital healthcare, SDG 9 because secure exchange is part of resilient digital infrastructure, and SDG 16 because signatures and audit trails strengthen accountability. If more resources were available, I would integrate HL7 FHIR, X.509 certificates, HSM-backed keys and tamper-evident audit storage. Overall, the work helped me attain CO2 by applying asymmetric-key architecture to a multi-party security problem and CO3 by evaluating hashes, digital signatures and access-control mechanisms for integrity, authentication and accountability.')

p('18. References')
refs=[
'Chandak, A., Chandak, P., & Soni, N. (2026). Blockchain applications in electronic health records: A systematic review of qualitative and quantitative evidence. BMC Medical Informatics and Decision Making, 26, 176. https://doi.org/10.1186/s12911-026-03476-3',
'Ullah, F., He, J., Zhu, N., Wajahat, A., Nazir, A., Qureshi, S., Pathan, M. S., & Dev, S. (2024). Blockchain-enabled EHR access auditing: Enhancing healthcare data security. Heliyon, 10(16), e34407. https://doi.org/10.1016/j.heliyon.2024.e34407',
'Tawfik, A. M., Al-Ahwal, A., Tag Eldien, A. S., et al. (2025). Blockchain-based access control and privacy preservation in healthcare: A comprehensive survey. Cluster Computing, 28, 529. https://doi.org/10.1007/s10586-025-05308-x',
'Thirasak, K., Chainarong, D., Chuaphanngam, T., et al. (2025). SSX-EHRs: Secure and scalable cross-domain EHRs sharing with blockchain sharding and dynamic proxy re-encryption. Journal on Information Security, 2025, 15. https://doi.org/10.1186/s13635-025-00200-y',
'Rajmohan, R., Suresh Kumar, K., Chowdhury, S., Sharma, B., et al. (2025). Blockchain-enabled authenticated key management framework for electronic health record systems. In Proceedings of the 1st International Conference on Creativity, Technology, and Sustainability, 97–108. https://doi.org/10.1007/978-981-97-8588-9_10',
'Tawfik, A. M., Al-Ahwal, A., Tag Eldien, A. S., et al. (2025). ACHealthChain blockchain framework for access control and privacy preservation in healthcare. Scientific Reports, 15, 16696. https://doi.org/10.1038/s41598-025-00757-1',
'Li et al. (2024). ECDSA-based tamper detection in medical data using a watermarking technique. International Journal of Cognitive Computing in Engineering, 5, 78–87. https://doi.org/10.1016/j.ijcce.2024.01.003',
'Fang et al. (2024). A blockchain-based hybrid encryption technique with anti-quantum signature for securing electronic health records. Complex & Intelligent Systems, 10, 6117–6141. https://doi.org/10.1007/s40747-024-01477-1',
'Chen, L., Moody, D., Regenscheid, A., Robinson, A. (2023). Digital Signature Standard (DSS), FIPS 186-5. National Institute of Standards and Technology. https://doi.org/10.6028/NIST.FIPS.186-5',
'National Institute of Standards and Technology. (2015, updated 2024). Secure Hash Standard (FIPS 180-4). https://doi.org/10.6028/NIST.FIPS.180-4',
'National Institute of Standards and Technology. (2015, planning update 2025). SHA-3 Standard: Permutation-Based Hash and Extendable-Output Functions (FIPS 202). https://doi.org/10.6028/NIST.FIPS.202',
'Josefsson, S., & Liusvaara, I. (2017). Edwards-Curve Digital Signature Algorithm (EdDSA). RFC 8032. RFC Editor.']
for r in refs: doc.add_paragraph(r,style='List Number')

p('Appendix A – Reproducibility and GitHub Checklist')
table(['Item','Included in ZIP','Verification'],[
('Source code','Yes','app/ + static/'),('Requirements','Yes','requirements.txt'),('Automated tests','Yes','pytest: 4 passed'),('Synthetic data','Yes','demo record in API'),('Screenshots','Yes','screenshots/'),('Figures','Yes','docs/'),('Report','Yes','DOCX + validation PDF'),('README','Yes','run instructions + feature list')])
p('Recommended local execution: create a virtual environment, install requirements, run `uvicorn app.main:app --reload`, open the displayed localhost address, and execute the Integrity Lab workflow. Run `pytest -q` from the project root to reproduce the automated tests.')

# footer/page numbers
for section in doc.sections:
    footer=section.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=footer.add_run('Assignment 5 • EHR CryptoGuard • '); r.font.size=Pt(8); r.font.color.rgb=RGBColor(100,100,100)
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); footer._p.append(fld)

out=ROOT/'Assignment_5_EHR_CryptoGuard_Final_Report.docx'
doc.save(out)
print(out)
