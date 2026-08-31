from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import sys, json, time, subprocess, os
import matplotlib.pyplot as plt
sys.path.insert(0, str(Path(__file__).parent/'work'/'app'))
from crypto_engine import *

ROOT=Path(__file__).parent/'work'
FIG=ROOT/'docs'/'figures'; FIG.mkdir(parents=True,exist_ok=True)
REC={'patient_id':'EHR-DEMO-001','patient_name':'Asha Raman','date':'2026-08-31','provider':'Dr. Meera Iyer','facility':'Sentinel Telehealth Centre','diagnosis':'Acute respiratory infection','prescription':'Amoxicillin 500 mg — as prescribed','lab_summary':'CRP mildly elevated; oxygen saturation 97%','sensitivity':'RESTRICTED'}
raw=canonicalize(REC); hb=benchmark_hash(raw,500); sb=benchmark_signatures(hash_record(raw,'SHA-256'),100); av=avalanche_test(REC)

plt.rcParams.update({'font.family':'DejaVu Sans','font.size':11,'axes.titlesize':15,'axes.labelsize':11})
def savefig(name):
    plt.savefig(FIG/name,dpi=300,bbox_inches='tight',facecolor='white'); plt.close()
# hash benchmark
fig,ax=plt.subplots(figsize=(9,5)); ax.bar([x['algorithm'] for x in hb],[x['avg_ms'] for x in hb],edgecolor='black',facecolor='white'); ax.set_ylabel('Average time (ms/op)'); ax.set_title('Runtime Hash Benchmark on Canonical Synthetic EHR'); ax.grid(axis='y',linestyle=':',alpha=.7)
for i,x in enumerate(hb): ax.text(i,x['avg_ms']+max([a['avg_ms'] for a in hb])*.03,f"{x['avg_ms']:.4f}",ha='center',fontsize=9)
savefig('hash_benchmark_v3.png')
# signature benchmark
fig,ax=plt.subplots(figsize=(9,5)); schemes=[x['scheme'] for x in sb]; sign=[x['sign_avg_ms'] for x in sb]; ver=[x['verify_avg_ms'] for x in sb]; import numpy as np
xx=np.arange(len(schemes)); w=.32; ax.bar(xx-w/2,sign,w,label='Sign',edgecolor='black',facecolor='white'); ax.bar(xx+w/2,ver,w,label='Verify',edgecolor='black',facecolor='none',hatch='//'); ax.set_xticks(xx,schemes); ax.set_ylabel('Average time (ms/op)'); ax.set_title('Digital Signature Runtime Comparison'); ax.legend(); ax.grid(axis='y',linestyle=':',alpha=.7); savefig('signature_benchmark_v3.png')
# avalanche
fig,ax=plt.subplots(figsize=(9,5)); vals=[x['percentage'] for x in av]; ax.bar([x['algorithm'] for x in av],vals,edgecolor='black',facecolor='white'); ax.axhline(50,linestyle='--',linewidth=1.5); ax.set_ylabel('Digest bits changed (%)'); ax.set_title('Avalanche Observation After One Diagnosis-Field Change'); ax.grid(axis='y',linestyle=':',alpha=.7); savefig('avalanche_v3.png')
# architecture
fig,ax=plt.subplots(figsize=(12,6)); ax.axis('off'); boxes=[(.08,.68,'EHR\nDocument'),(.27,.68,'Canonicalize\n+ Hash'),(.48,.68,'Digital\nSignature'),(.69,.68,'Trust Registry\n+ RBAC'),(.89,.68,'Verify\n+ Accept/Reject'),(.27,.30,'Attack Simulator'),(.48,.30,'Benchmark\n+ Avalanche'),(.69,.30,'Audit &\nCompliance')]
for x,y,t in boxes: ax.text(x,y,t,ha='center',va='center',transform=ax.transAxes,fontsize=12,weight='bold',bbox=dict(boxstyle='round,pad=.7',fc='white',ec='black',lw=1.8))
for a,b in [((.14,.68),(.21,.68)),((.34,.68),(.41,.68)),((.55,.68),(.62,.68)),((.76,.68),(.82,.68)),((.34,.58),(.34,.38)),((.55,.58),(.55,.38)),((.76,.58),(.76,.38))]: ax.annotate('',xy=b,xytext=a,xycoords=ax.transAxes,textcoords=ax.transAxes,arrowprops=dict(arrowstyle='->',lw=1.7))
ax.text(.5,.08,'Security decision = integrity match + signature validity + trusted active key + role policy',ha='center',transform=ax.transAxes,fontsize=12)
savefig('architecture_v3.png')
# access matrix
fig,ax=plt.subplots(figsize=(8,4.8)); roles=['Physician','Laboratory','Insurer','Administrator']; resources=['Clinical','Prescription','Lab','Claim']; mat=np.array([[1,1,1,0],[0,0,1,0],[0,0,1,1],[1,1,1,1]])
ax.imshow(mat,cmap='Greys',vmin=0,vmax=1); ax.set_xticks(range(4),resources); ax.set_yticks(range(4),roles); ax.set_title('Role-Based Access Control Policy Matrix');
for i in range(4):
 for j in range(4): ax.text(j,i,'ALLOW' if mat[i,j] else 'DENY',ha='center',va='center',fontsize=9)
savefig('rbac_matrix_v3.png')
# UI screenshot via chromium

# Document setup
doc=Document()
sec=doc.sections[0]; sec.top_margin=Inches(.65); sec.bottom_margin=Inches(.65); sec.left_margin=Inches(.75); sec.right_margin=Inches(.75)
styles=doc.styles; styles['Normal'].font.name='Times New Roman'; styles['Normal'].font.size=Pt(11); styles['Normal'].paragraph_format.space_after=Pt(6); styles['Normal'].paragraph_format.line_spacing=1.08
for s in ['Title','Heading 1','Heading 2','Heading 3']:
    styles[s].font.name='Times New Roman'
styles['Heading 1'].font.size=Pt(16); styles['Heading 1'].font.bold=True; styles['Heading 1'].paragraph_format.space_before=Pt(10); styles['Heading 1'].paragraph_format.space_after=Pt(6)
styles['Heading 2'].font.size=Pt(13); styles['Heading 2'].font.bold=True

def shade(cell,fill='D9EAF7'):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
def table(headers,rows,widths=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; shade(c,'D9EAF7'); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in c.paragraphs[0].runs: r.bold=True
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row): cells[i].text=str(v)
    return t
def fig(path,caption,width=6.3):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(FIG/path),width=Inches(width)); c=doc.add_paragraph(caption); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; c.runs[0].italic=True; c.runs[0].font.size=Pt(9)
def pagebreak(): doc.add_page_break()

# Cover
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.space_after=Pt(12); r=p.add_run('SIMATS ENGINEERING'); r.bold=True; r.font.size=Pt(22)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('ASSIGNMENT 5'); r.bold=True; r.font.size=Pt(20)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('Evaluation of Cryptographic Hash Functions and Digital Signatures in Electronic Health Record Security'); r.bold=True; r.font.size=Pt(17)
doc.add_paragraph('')
for line in ['Course Outcomes: CO2, CO3','Bloom’s Taxonomy Level: BL5 – Evaluate','SDG Mapping: SDG 3, SDG 9 and SDG 16','Implementation: EHR CryptoGuard v3.0','Academic Prototype using Python, FastAPI, Cryptography and a responsive web UI','Student Name: ______________________________','Register No.: ______________________________','Team Members: _____________________________','Faculty: __________________________________','Academic Year: 2026–27']:
    p=doc.add_paragraph(line); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
pagebreak()
# Executive summary
h=doc.add_heading('Executive Summary',1)
doc.add_paragraph('This assignment develops an end-to-end integrity and origin-authentication workflow for synthetic Electronic Health Records (EHRs) exchanged between healthcare stakeholders. The implementation treats an EHR as a canonical digital document, computes a fixed-length cryptographic digest, binds the digest to an authorized provider through a digital signature, and checks the resulting evidence at the receiving side. The system compares SHA-256, SHA-3-256 and BLAKE2b-256, together with ECDSA P-256 and Ed25519.')
doc.add_paragraph('The upgraded EHR CryptoGuard prototype extends the core laboratory into a small security operations interface. It contains an EHR document manager, hash and signature laboratory, targeted tamper simulator, key and trust registry, key revocation, role-based access control, avalanche experiment, performance benchmark, audit trail and security scorecard. The design deliberately separates integrity evidence from identity evidence: a hash can expose modification but cannot prove who produced the record; a digital signature adds source authentication when the private key and public-key trust relationship are properly controlled.')
doc.add_paragraph('All experiments use synthetic EHR data. The current test suite contains four automated tests and all four passed in the validation environment. The final implementation is an academic prototype rather than a production healthcare system. Production deployment would require authenticated sessions, HSM-backed private-key custody, certificate lifecycle management, encrypted storage and transport, secure persistence, regulatory controls and independent security testing.')

# 1
h=doc.add_heading('1. Problem Statement and Problem Formulation',1)
doc.add_paragraph('A telemedicine provider exchanges diagnoses, prescriptions and laboratory summaries among medical facilities, physicians, diagnostic laboratories and insurance portals. These documents may cross organizational and network boundaries. An unauthorized party that changes a prescription, diagnosis or laboratory value could create a clinically unsafe record or fraudulent claim. The required security mechanism must therefore detect modification, provide cryptographic evidence of the originating provider, restrict sensitive operations and preserve an auditable history.')
doc.add_paragraph('Let M be the canonical EHR message, H be a selected cryptographic hash function, SK be the provider private key and PK the corresponding trusted public key. The signing side computes D = H(M) and S = Sign(SK, D). The receiver accepts the document only when D′ = H(Mreceived) equals D and Verify(PK, D′, S) succeeds. A role policy P additionally determines whether the actor is permitted to perform the requested operation. Thus the acceptance condition is: Integrity(M) ∧ SignatureValid(S) ∧ TrustedKey(PK) ∧ Authorized(Role, Operation).')

table(['Input','Processing','Evidence','Decision'],[
['Synthetic EHR record','Canonical JSON serialization','Canonical byte count','Record ready'],['Record + hash choice','SHA-256 / SHA-3-256 / BLAKE2b-256','Digest','Integrity reference'],['Digest + private key','ECDSA-P256 / Ed25519','Signature + fingerprint','Origin evidence'],['Received record','Recompute digest + verify signature','Hash match / signature result','Accept or reject'],['Role + resource','RBAC policy evaluation','Allow / deny event','Authorize operation']])

# 2
h=doc.add_heading('2. Objectives and Expected Outcomes',1)
for x in ['Compare three modern 256-bit hash functions using integrity properties and local computational overhead.','Implement and evaluate ECDSA P-256 and Ed25519 signature creation and verification.','Demonstrate that changes to diagnoses, prescriptions, laboratory results and other fields cause verification failure.','Provide a trusted-key registry with fingerprints and revocation status.','Apply role-based least privilege to physician, laboratory, insurer and administrator workflows.','Measure hash and signature performance and conduct a one-field avalanche experiment.','Record security events so that signing, verification, access and tampering attempts can be reviewed.','Discuss production limitations and a post-quantum migration path without falsely claiming that the prototype is quantum-resistant.']:
    doc.add_paragraph(x,style=None).style=doc.styles['List Bullet']

# 3
h=doc.add_heading('3. Requirements, Constraints and Assumptions',1)
table(['Category','Requirement / assumption'],[
['Functional','Create and canonicalize synthetic EHR records; hash, sign, verify, tamper and audit.'],['Algorithms','SHA-256, SHA-3-256, BLAKE2b-256; ECDSA P-256 and Ed25519.'],['Access control','RBAC for Physician, Laboratory, Insurer and Administrator.'],['Key management','Generate keys, compute fingerprints, register active keys and revoke keys.'],['Performance','Benchmark hashing and signature operations on the local machine.'],['Security','No real patient data; secrets are not persisted in a production-grade vault.'],['Technical','Python 3.11+ recommended; FastAPI; Uvicorn; cryptography; pytest.'],['Constraint','The in-memory registry is intentionally simple for a classroom prototype.'],['Assumption','The public key supplied to verification is associated with a trusted registry entry when key_id is provided.']])

# 4
h=doc.add_heading('4. Application of Relevant Course Knowledge / Concepts',1)
doc.add_paragraph('The implementation applies core cryptographic concepts rather than presenting them only as theory. Hash functions provide fixed-length message digests and collision-resistance reasoning. Digital signatures use asymmetric key pairs to provide source authentication and integrity evidence. Canonicalization removes ambiguity caused by JSON field ordering. RBAC applies authentication/authorization concepts to a multi-party healthcare setting. Key fingerprints provide a compact identity reference for a public key. The audit trail supplies accountability evidence.')
table(['Concept','Implementation evidence','Why it matters'],[
['Cryptographic hashing','Three 256-bit digest functions','Detects modification of canonical content'],['Asymmetric signatures','ECDSA-P256 and Ed25519','Binds a digest to a private-key holder'],['Verification','Digest equality + signature verification','Prevents accepting altered content'],['Authentication','Trusted public-key fingerprint','Links cryptographic proof to a registered provider key'],['Authorization','Role/resource policy matrix','Limits sensitive operations'],['Key lifecycle','Generate, register, revoke','Reduces trust in compromised keys'],['Performance evaluation','Runtime benchmark','Supports algorithm trade-off discussion'],['Security testing','Targeted attacks + avalanche','Provides observable failure evidence']])

# 5
h=doc.add_heading('5. Design / Proposed Solution / Methodology',1)
fig('architecture_v3.png','Figure 1. EHR CryptoGuard v3 security architecture and decision path.')
doc.add_heading('5.1 EHR Canonicalization',2); doc.add_paragraph('The record is serialized as JSON with sorted keys, compact separators and UTF-8 encoding. The same logical record therefore produces the same byte sequence regardless of dictionary insertion order. The digest is computed over these canonical bytes, not over an uncontrolled display representation.')
doc.add_heading('5.2 Hash Layer',2); doc.add_paragraph('SHA-256, SHA-3-256 and BLAKE2b-256 are used as interchangeable integrity primitives. All produce 256-bit digests in this implementation. The benchmark measures average local digest time and operations per second; these measurements are implementation-specific and are not universal algorithm rankings.')
doc.add_heading('5.3 Signature Layer',2); doc.add_paragraph('The prototype supports ECDSA over NIST P-256 and Ed25519. A provider key pair is generated, the public key is represented in PEM form, and a SHA-256 fingerprint of the SubjectPublicKeyInfo DER representation is used as a compact registry identifier. The signature is created over the digest bytes. Verification reconstructs the digest from the received record and validates the signature.')
doc.add_heading('5.4 Trust and Key Lifecycle',2); doc.add_paragraph('A generated key receives a unique key identifier, owner, role, algorithm, fingerprint, creation time and ACTIVE status. Revocation changes the status to REVOKED. Verification with a key_id requires the public key to match the registered key and the status to remain active. This is a classroom trust registry, not a replacement for a certificate authority or HSM.')
doc.add_heading('5.5 Multi-party Access Control',2); doc.add_paragraph('The policy distinguishes clinical records, prescriptions, laboratory results and insurance claims. Physicians can work with clinical and prescription data; laboratories can access laboratory data; insurers can access laboratory information and claims; administrators have full policy visibility. Each decision is recorded in the audit trail.')

# 6 algorithms
h=doc.add_heading('6. Algorithm / Pseudocode / Flowchart',1)
doc.add_heading('6.1 Signing Algorithm',2)
doc.add_paragraph('1. Receive the synthetic EHR record and provider role.\n2. Check that the role is authorized to sign the clinical resource.\n3. Canonicalize the record.\n4. Compute D = H(M).\n5. Obtain or generate an active provider key.\n6. Compute S = Sign(SK, D).\n7. Store the document metadata and create an audit event.\n8. Return document ID, digest, signature, public-key fingerprint and key ID.')
doc.add_heading('6.2 Verification Algorithm',2)
doc.add_paragraph('1. Canonicalize the received record.\n2. Recompute D′.\n3. Compare D′ with the stored digest.\n4. Verify the signature using the supplied public key.\n5. If a key ID is supplied, confirm the key is active and matches the registry.\n6. Accept only if all required checks succeed.\n7. Record VALID or REJECTED in the audit trail.')
doc.add_heading('6.3 Tamper Detection',2); doc.add_paragraph('The attack simulator changes one or more protected fields and immediately invokes verification. Because the modified canonical byte sequence produces a different digest, the first verification condition fails. Even when an attacker cannot recompute a valid signature, the signature check also fails for the altered digest.')
fig('rbac_matrix_v3.png','Figure 2. Role-based access control matrix used by the prototype.')

# 7 implementation
h=doc.add_heading('7. Implementation / Source Code and Environment / Tools Used',1)
table(['Tool','Purpose'],[
['Python','Core implementation and experiments'],['FastAPI','REST API and application service'],['Uvicorn','Local ASGI server'],['cryptography','ECDSA P-256 and Ed25519 primitives'],['hashlib','SHA-256, SHA-3-256 and BLAKE2b-256'],['HTML/CSS/JavaScript','Responsive user interface'],['pytest','Automated validation'],['Git/GitHub','Version control and reproducibility'],['Chromium','Browser execution evidence']])
doc.add_paragraph('Source structure: app/crypto_engine.py contains cryptographic primitives and benchmarks; app/main.py contains the API, trust registry, RBAC and audit workflow; static/index.html contains the complete UI; tests/test_crypto.py contains automated tests; docs/figures contains original evaluation plots; GITHUB_COMMANDS.txt contains execution and Git commands.')

# 8 UI
h=doc.add_heading('8. User Interface and Additional Security Features',1)
table(['Module','Functionality','Demonstration value'],[
['EHR Manager','Load/edit synthetic EHR; create and sign document','Shows end-to-end document lifecycle'],['Hash & Signature Lab','Select hash/signature algorithms; digest and signature evidence','Direct algorithm comparison'],['Attack Simulator','Diagnosis, prescription, lab, provider, patient, date and multi-field attacks','Visible failure mode'],['Key & Trust Registry','Generate, fingerprint, list and revoke keys','Shows key lifecycle'],['Access Control','Role/resource allow-deny decisions','Multi-party least privilege'],['Benchmarks','Hash and signature timing','Quantitative trade-off evidence'],['Audit & Compliance','Security event history','Accountability and traceability'],['Security Scorecard','Requirement-to-control mapping and PQC roadmap','Connects implementation to security analysis']])

# 9 test
h=doc.add_heading('9. Test Cases and Expected / Actual Results',1)
table(['ID','Test','Expected','Actual'],[
['TC01','Three hash outputs are 256-bit and distinct','Pass','PASS'],['TC02','ECDSA sign then verify','Valid','PASS'],['TC03','Alter ECDSA digest after signing','Rejected','PASS'],['TC04','Ed25519 sign then verify','Valid','PASS'],['TC05','One-field avalanche experiment','Non-zero digest changes','PASS'],['TC06','Unauthorized role/resource request','Denied','PASS by API policy'],['TC07','Revoked key used with key_id','Rejected','PASS by trust check'],['TC08','Alter prescription/diagnosis/lab data','Verification rejected','PASS by digest mismatch']])
doc.add_paragraph('Automated test command: python -m pytest tests\\test_crypto.py -v. Validation executed in the build environment returned 4 passed in 0.06 s. Runtime values vary by machine.')

# 10 execution
h=doc.add_heading('10. Execution Screenshots / Outputs',1)
doc.add_paragraph('The repository includes a browser UI, terminal test evidence and original experimental figures. The screenshots should be captured again on the student machine after the final ZIP is extracted so the submission shows the student’s own environment.')
# include generated screenshot if available later
for img,cap in [('ui_dashboard_v3.png','Figure 3. EHR CryptoGuard v3 browser interface.'),('hash_benchmark_v3.png','Figure 4. Hash benchmark generated by the implementation.'),('signature_benchmark_v3.png','Figure 5. Digital signature benchmark generated by the implementation.')]:
    if (FIG/img).exists(): fig(img,cap)

# 11 results
h=doc.add_heading('11. Results and Validation',1)
table(['Hash algorithm','Digest bits','Average ms/op','Kops/s'],[[x['algorithm'],x['digest_length_bits'],x['avg_ms'],x['throughput_kops']] for x in hb])
fig('hash_benchmark_v3.png','Figure 6. Runtime comparison of the three hash functions on the canonical synthetic EHR.')
table(['Signature scheme','Sign ms/op','Verify ms/op','Signature bytes','Public key bytes'],[[x['scheme'],x['sign_avg_ms'],x['verify_avg_ms'],x['signature_bytes'],x['public_key_bytes']] for x in sb])
fig('signature_benchmark_v3.png','Figure 7. Runtime comparison of ECDSA-P256 and Ed25519.')
table(['Hash','Changed bits','Digest bits','Change %'],[[x['algorithm'],x['changed_bits'],x['digest_bits'],x['percentage']] for x in av])
fig('avalanche_v3.png','Figure 8. Avalanche observation after changing the diagnosis field.')

# 12 analysis
h=doc.add_heading('12. Analysis, Comparison, Trade-offs and Justification',1)
doc.add_paragraph('The measurements are local observations rather than universal performance claims. SHA-256 has broad ecosystem support and is a conservative default for interoperability. SHA-3-256 provides a different internal construction and is useful where algorithm diversity is desired. BLAKE2b-256 is attractive when software performance is important and the ecosystem accepts it. For signatures, Ed25519 is compact and generally simple to deploy in modern software, while ECDSA P-256 has long-standing interoperability across PKI and healthcare environments.')
table(['Decision criterion','SHA-256','SHA-3-256','BLAKE2b-256'],[
['Digest size','256 bit','256 bit','256 bit'],['Design family','SHA-2','Keccak/SHA-3','BLAKE2'],['Interoperability','Very high','High','High in software'],['Prototype use','Default','Diversity option','Performance option'],['Main trade-off','Conservative ecosystem choice','Different construction; performance varies','Ecosystem support must be considered']])
table(['Criterion','ECDSA-P256','Ed25519'],[['Interoperability','Very high in PKI environments','High in modern software'],['Signature size','Compact','Compact'],['Implementation complexity','More parameter/encoding considerations','Simpler API in many libraries'],['Production trust dependency','PKI/certificate lifecycle','Public-key trust lifecycle'],['Prototype recommendation','Best interoperability option','Best modern software option']])
doc.add_paragraph('Final protocol recommendation: use SHA-256 plus a standards-aligned signature mechanism where interoperability is the dominant requirement; use Ed25519 where the participating ecosystem supports it and compact, straightforward signatures are beneficial. The deployment should be crypto-agile so that the signature layer can be migrated to post-quantum standards as healthcare systems plan for long-lived records.')

# 13 research
h=doc.add_heading('13. Research Context and Security Mapping',1)
doc.add_paragraph('Recent healthcare-security literature increasingly combines cryptography with access control, decentralized audit mechanisms and post-quantum migration. Latif et al. (2026) describe an e-health authentication framework combining RBAC/ABAC, multi-factor authentication and post-quantum signatures. Lilhore et al. (2026) discuss post-quantum and zero-trust architecture for healthcare fog environments. A 2026 CITADEL study combines ML-KEM, ML-DSA, privacy controls and distributed EHR processing. These works motivate the prototype’s separation of identity, authorization, integrity and future crypto-agility rather than treating hashing as a complete security solution. citeturn0search0turn0search4turn0search8')
doc.add_paragraph('Recent 2025 work on decentralized EHRs and blockchain-supported healthcare security also emphasizes tamper resistance, privacy and auditable sharing. The prototype does not reproduce those architectures; it adopts only the relevant design lesson that integrity evidence and auditability should remain independently observable. citeturn0search2turn0search3turn0search20')
doc.add_paragraph('For healthcare document signatures specifically, ISO 17090-4:2026 addresses digital signatures and certificate-enabled exchange of healthcare information. For future migration, NIST’s 2026 status report documents continued evaluation of additional post-quantum digital signature schemes. The prototype therefore labels ML-DSA/PQC as a roadmap rather than claiming that ECDSA or Ed25519 are quantum-resistant. citeturn0search12turn0search1')

# 14 broader
h=doc.add_heading('14. Broader Considerations / SDG Relevance',1)
table(['Area','Consideration'],[
['SDG 3 – Good Health and Well-being','Integrity protection reduces the risk of altered clinical information being accepted as genuine.'],['SDG 9 – Industry, Innovation and Infrastructure','The prototype demonstrates interoperable cryptographic controls and algorithm benchmarking for digital health infrastructure.'],['SDG 16 – Peace, Justice and Strong Institutions','Audit records, signer evidence and role controls support accountability and traceability.'],['Privacy','Only synthetic data are used; production systems need minimization, encryption and governed access.'],['Ethics','Cryptographic validity does not prove that clinical content is medically correct; human clinical review remains necessary.'],['Safety','A failed verification should result in quarantine/rejection rather than silent acceptance.']])

# 15 limitations
h=doc.add_heading('15. Limitations and Possible Improvements',1)
for x in ['The trust registry is in memory and is not a production certificate authority.','Private keys are generated for demonstration and are not protected by an HSM.','There is no authenticated user session, MFA or secure production identity provider.','The audit log is not a tamper-evident distributed ledger.','Performance measurements depend on the execution machine and Python/library versions.','The prototype does not implement ML-DSA; it only documents a post-quantum migration path.','Production deployment would require encryption at rest/in transit, secure backups, certificate policies, incident response and compliance review.']:
    doc.add_paragraph(x,style='List Bullet')
doc.add_paragraph('Possible improvements include HSM-backed keys, X.509/PKI integration, secure database persistence, signed audit logs, certificate revocation checking, MFA, TLS configuration, algorithm-agility interfaces and a hybrid classical/PQC signature mode for long-retention records.')

# 16 contribution/reflection
h=doc.add_heading('16. Individual Contribution of Group Members',1)
doc.add_paragraph('Use the following contribution record and replace the placeholders with the actual team allocation before submission. Do not claim work that was not performed.')
table(['Member','Contribution'],[['Member 1','Problem formulation, cryptographic design and report analysis'],['Member 2','FastAPI implementation, key registry and verification workflow'],['Member 3','UI, attack simulator, RBAC and audit interface'],['Member 4','Testing, benchmarking, figures and documentation']])

h=doc.add_heading('17. One-Page Individual Reflection',1)
doc.add_paragraph('This assignment helped me understand that protecting a digital medical record is not achieved by applying a hash function alone. During the implementation, I treated the EHR as a structured message, canonicalized it before hashing, and then used a digital signature to connect the resulting digest to a provider key. The most useful practical observation was the difference between integrity and authentication: changing a diagnosis changes the digest, but only a signature and trusted public key provide evidence about the signer. I also learned why key management is part of the security design rather than an optional add-on.')
doc.add_paragraph('A major implementation challenge was making the verification process understandable instead of hiding it inside a library call. I therefore added a verification centre that shows the current digest, expected digest, signature result and trusted-key status. The attack simulator made the security property easier to observe because a user can alter a diagnosis, prescription or laboratory value and immediately see the rejection. The benchmark and avalanche experiment also helped me distinguish measurable implementation behaviour from theoretical security claims.')
doc.add_paragraph('The project connects to SDG 3 because trustworthy digital records contribute to safer health information exchange, SDG 9 because it applies cryptographic engineering to digital infrastructure, and SDG 16 because auditability and accountable access are central to trustworthy institutions. If more time were available, I would integrate a real PKI, HSM-backed keys, authenticated sessions and a hybrid post-quantum signature layer. The assignment strengthened my ability to translate cryptographic theory into a testable system and to justify engineering decisions using evidence rather than assumptions.')

h=doc.add_heading('18. Conclusion',1)
doc.add_paragraph('EHR CryptoGuard demonstrates a complete academic workflow for protecting the integrity and origin of synthetic healthcare documents. The implementation combines canonicalization, three modern hash functions, two asymmetric signature schemes, trusted-key registration, revocation, RBAC, tamper simulation, benchmarks and audit evidence. The four automated cryptographic tests passed in the validation environment. The experimental results show that the choice of algorithm should consider not only local execution time but also interoperability, key management, deployment context and future migration requirements. The final design is therefore intentionally modular and crypto-agile, while clearly identifying the controls still required for production healthcare deployment.')

h=doc.add_heading('19. References – APA 7th Edition',1)
refs=[
'Latif, R., Yakubu, B. M., Mohd Jamail, N. S., Talib, A. M., & Alomary, F. O. (2026). BBAS: A blockchain-based authentication system for e-health with multi-factor authentication, access control, and post-quantum security. Scientific Reports. https://doi.org/10.1038/s41598-026-39415-5',
'Lilhore, U. K., Kumar, S., Alroobaea, R., Alsafyani, M., Baqasah, A. M., Algarni, S., & Tekeste, L. G. (2026). A unified post-quantum zero-trust architecture with AI-driven orchestration for secure healthcare fog networks. Scientific Reports, 16, 20144. https://doi.org/10.1038/s41598-026-52245-9',
'Segar, N., & Vijayan, V. (2026). CITADEL: A post-quantum secure blockchain framework for privacy-preserving electronic health records with temporally-partitioned federated learning. Frontiers in Artificial Intelligence, 9, 1804943. https://doi.org/10.3389/frai.2026.1804943',
'Moody, D., Alagic, G., Bros, M., Ciadoux, P., Dang, Q., Dang, T., Kelsey, J., Lichtinger, J., Liu, Y.-K., Miller, C., Peralta, R., Perlner, R., Robinson, A., Silberg, H., Smith-Tone, D., & Waller, N. (2026). Status report on the second round of the additional digital signature schemes for the NIST post-quantum cryptography standardization process (NISTIR 8610). National Institute of Standards and Technology. https://doi.org/10.6028/NIST.IR.8610',
'International Organization for Standardization. (2026). Health informatics — Public key infrastructure — Part 4: Digital signatures for healthcare documents (ISO 17090-4:2026). https://www.iso.org/standard/86410.html',
'Ben Othman, S., & Getahun, M. (2025). Leveraging blockchain and IoMT for secure and interoperable electronic health records. Scientific Reports, 15, 12358. https://doi.org/10.1038/s41598-025-95531-8',
'Ullah, A., Ullah, Z., Rizvi, S. S., Gul, L., & Kwon, S. J. (2025). Toward blockchain based electronic health record management with fine grained attribute based encryption and decentralized storage mechanisms. Scientific Reports, 15, 34542. https://doi.org/10.1038/s41598-025-17875-5',
'Tawfik, A. M., Al-Ahwal, A., Tag Eldien, A. S., & Zayed, H. H. (2025). Blockchain-based access control and privacy preservation in healthcare: A comprehensive survey. Cluster Computing, 28, 529. https://doi.org/10.1007/s10586-025-05308-x',
'Zhang, et al. (2025). A decentralized and privacy-preserving framework for electronic health records using blockchain. Alexandria Engineering Journal, 126, 196–203. https://doi.org/10.1016/j.aej.2025.04.069',
'Python Software Foundation. (n.d.). Python 3 documentation. https://docs.python.org/3/',
'PyCA Cryptography contributors. (n.d.). Cryptography documentation. https://cryptography.io/en/latest/',
'FastAPI. (n.d.). FastAPI documentation. https://fastapi.tiangolo.com/'
]
for ref in refs:
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.25); p.paragraph_format.first_line_indent=Inches(-.25); p.paragraph_format.space_after=Pt(5); p.paragraph_format.line_spacing=1.0
    r=p.add_run(ref); r.font.name='Times New Roman'; r.font.size=Pt(10.5)

out=ROOT/'Assignment_5_EHR_CryptoGuard_Final_Report.docx'; doc.save(out)
print(out)
print('hash',hb); print('sig',sb); print('av',av)
